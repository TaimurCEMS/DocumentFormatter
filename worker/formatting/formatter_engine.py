"""Formatter engine for format-only DOCX processing."""

from docx import Document
from docx.shared import Cm
from .format_profiles import get_profile, FormatProfile
from .docx_utils import (
    set_section_margins,
    set_page_size,
    apply_normal_style,
    update_normal_style_definition,
    normalize_heading_styles,
    extract_plain_text,
    docx_to_bytes,
    bytes_to_docx
)


def apply_format_only(docx_bytes: bytes, profile_name: str = "standard_clean"):
    """Apply format-only changes to a DOCX document.
    
    This function:
    - Preserves all text content exactly
    - Only modifies formatting properties (margins, fonts, spacing, etc.)
    - Extracts plain text for storage
    
    Args:
        docx_bytes: Original DOCX file as bytes
        profile_name: Name of format profile to apply (default: "standard_clean")
        
    Returns:
        tuple: (formatted_docx_bytes, extracted_text)
            - formatted_docx_bytes: Formatted DOCX as bytes
            - extracted_text: Plain text extracted from document
    """
    # Get formatting profile
    profile = get_profile(profile_name)
    
    # Load document
    doc = bytes_to_docx(docx_bytes)
    
    # Extract text BEFORE formatting (to ensure we preserve content)
    extracted_text = extract_plain_text(doc)
    
    # Step 1: Update Normal style definition in styles.xml
    update_normal_style_definition(
        doc,
        profile.normal_font_name,
        profile.normal_font_size,
        profile.line_spacing,
        profile.paragraph_spacing_before,
        profile.paragraph_spacing_after
    )
    
    # Step 2: Apply formatting to all sections (margins and page size)
    for section in doc.sections:
        # Set page size
        set_page_size(section, profile.page_width, profile.page_height)
        
        # Set margins directly using Cm() - convert cm floats to Cm() at assignment time
        # This ensures exact Cm(2.0) = 720000 EMU, not Twips conversion
        section.top_margin = Cm(profile.margins['top'])
        section.bottom_margin = Cm(profile.margins['bottom'])
        section.left_margin = Cm(profile.margins['left'])
        section.right_margin = Cm(profile.margins['right'])
    
    # Step 3: Force-apply Normal style formatting to ALL paragraphs
    # (Many docs override Normal and won't inherit, so we must apply directly)
    for paragraph in doc.paragraphs:
        apply_normal_style(
            paragraph,
            profile.normal_font_name,
            profile.normal_font_size,
            profile.line_spacing,
            profile.paragraph_spacing_before,
            profile.paragraph_spacing_after
        )
    
    # Step 4: Normalize heading styles (only if already using heading styles)
    normalize_heading_styles(doc)
    
    # Step 5: Force-apply formatting to ALL table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_normal_style(
                        paragraph,
                        profile.normal_font_name,
                        profile.normal_font_size,
                        profile.line_spacing,
                        profile.paragraph_spacing_before,
                        profile.paragraph_spacing_after
                    )
    
    # Convert back to bytes
    formatted_bytes = docx_to_bytes(doc)
    
    # Verify text content is unchanged
    doc_after = bytes_to_docx(formatted_bytes)
    extracted_text_after = extract_plain_text(doc_after)
    
    if extracted_text != extracted_text_after:
        # This should never happen, but log if it does
        print(f"WARNING: Text content changed during format-only processing!")
        print(f"Before: {len(extracted_text)} chars")
        print(f"After: {len(extracted_text_after)} chars")
    
    return formatted_bytes, extracted_text

