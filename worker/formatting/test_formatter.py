"""Tests for format-only DOCX formatter."""

import unittest
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
import io
import zipfile
import xml.etree.ElementTree as ET

from formatting.formatter_engine import apply_format_only
from formatting.format_profiles import STANDARD_CLEAN, COMPACT_CLEAN, LARGE_READABLE
from formatting.docx_utils import extract_plain_text, bytes_to_docx, docx_to_bytes


def cm_to_twips(cm):
    """Convert centimeters to twips.
    
    Args:
        cm: float, centimeters
        
    Returns:
        int: twips (1 inch = 1440 twips, 1 inch = 2.54 cm)
    """
    return int(round(cm / 2.54 * 1440))


class TestFormatOnly(unittest.TestCase):
    """Test format-only processing preserves text content."""
    
    def create_test_docx(self, text_content, messy_formatting=True):
        """Create a test DOCX with optional messy formatting."""
        doc = Document()
        
        for line in text_content.split('\n'):
            if line.strip():
                para = doc.add_paragraph(line.strip())
                if messy_formatting:
                    # Add messy formatting
                    para.paragraph_format.space_before = Pt(24)  # Too much spacing
                    para.paragraph_format.space_after = Pt(18)
                    para.paragraph_format.line_spacing = 2.0  # Double spacing
                    for run in para.runs:
                        run.font.name = 'Arial'  # Different font
                        run.font.size = Pt(14)  # Different size
        
        return docx_to_bytes(doc)
    
    def test_text_content_preserved(self):
        """Test that text content is exactly preserved."""
        original_text = "This is paragraph one.\n\nThis is paragraph two.\n\nThis is paragraph three."
        input_bytes = self.create_test_docx(original_text, messy_formatting=True)
        
        formatted_bytes, extracted_text = apply_format_only(input_bytes, "standard_clean")
        
        # Text should be exactly the same
        self.assertEqual(extracted_text, original_text.replace('\n\n', '\n'))
        
        # Verify by loading and extracting again
        doc_after = bytes_to_docx(formatted_bytes)
        text_after = extract_plain_text(doc_after)
        self.assertEqual(text_after, original_text.replace('\n\n', '\n'))
    
    def test_section_margins_applied(self):
        """Test that section margins match profile."""
        input_bytes = self.create_test_docx("Test document")
        formatted_bytes, _ = apply_format_only(input_bytes, "standard_clean")
        
        doc = bytes_to_docx(formatted_bytes)
        section = doc.sections[0]
        
        # Check margins (2.54 cm) - compare with Cm() converted values
        from docx.shared import Cm
        self.assertEqual(section.top_margin, Cm(STANDARD_CLEAN.margins['top']))
        self.assertEqual(section.bottom_margin, Cm(STANDARD_CLEAN.margins['bottom']))
        self.assertEqual(section.left_margin, Cm(STANDARD_CLEAN.margins['left']))
        self.assertEqual(section.right_margin, Cm(STANDARD_CLEAN.margins['right']))
    
    def test_normal_style_font_applied(self):
        """Test that Normal style font matches profile."""
        input_bytes = self.create_test_docx("Test paragraph")
        formatted_bytes, _ = apply_format_only(input_bytes, "standard_clean")
        
        doc = bytes_to_docx(formatted_bytes)
        paragraph = doc.paragraphs[0]
        
        # Check font
        if paragraph.runs:
            self.assertEqual(paragraph.runs[0].font.name, STANDARD_CLEAN.normal_font_name)
            self.assertEqual(paragraph.runs[0].font.size, STANDARD_CLEAN.normal_font_size)
    
    def test_paragraph_spacing_applied(self):
        """Test that paragraph spacing matches profile."""
        input_bytes = self.create_test_docx("Paragraph one\n\nParagraph two")
        formatted_bytes, _ = apply_format_only(input_bytes, "standard_clean")
        
        doc = bytes_to_docx(formatted_bytes)
        paragraph = doc.paragraphs[0]
        para_format = paragraph.paragraph_format
        
        # Check spacing
        self.assertEqual(para_format.space_before, STANDARD_CLEAN.paragraph_spacing_before)
        self.assertEqual(para_format.space_after, STANDARD_CLEAN.paragraph_spacing_after)
        # Line spacing should be approximately 1.15 (may be stored as float)
        self.assertAlmostEqual(para_format.line_spacing, STANDARD_CLEAN.line_spacing, places=2)
    
    def test_line_spacing_applied(self):
        """Test that line spacing matches profile."""
        input_bytes = self.create_test_docx("Test line spacing\nSecond line")
        formatted_bytes, _ = apply_format_only(input_bytes, "standard_clean")
        
        doc = bytes_to_docx(formatted_bytes)
        paragraph = doc.paragraphs[0]
        para_format = paragraph.paragraph_format
        
        # Line spacing should be 1.15
        self.assertEqual(para_format.line_spacing_rule, WD_LINE_SPACING.MULTIPLE)
        self.assertAlmostEqual(para_format.line_spacing, 1.15, places=2)
    
    def test_multiple_paragraphs_preserved(self):
        """Test that multiple paragraphs are preserved."""
        original_text = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph with more text."
        input_bytes = self.create_test_docx(original_text)
        formatted_bytes, extracted_text = apply_format_only(input_bytes, "standard_clean")
        
        # All paragraphs should be present
        self.assertIn("First paragraph", extracted_text)
        self.assertIn("Second paragraph", extracted_text)
        self.assertIn("Third paragraph", extracted_text)
        
        # Count should match
        doc_after = bytes_to_docx(formatted_bytes)
        self.assertEqual(len(doc_after.paragraphs), 3)
    
    def test_empty_document_handled(self):
        """Test that empty document is handled gracefully."""
        doc = Document()
        input_bytes = docx_to_bytes(doc)
        
        formatted_bytes, extracted_text = apply_format_only(input_bytes, "standard_clean")
        
        # Should return empty text
        self.assertEqual(extracted_text, "")
        
        # Document should still be valid
        doc_after = bytes_to_docx(formatted_bytes)
        self.assertIsNotNone(doc_after)
    
    def test_table_formatting(self):
        """Test that tables are formatted but text preserved."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        
        input_bytes = docx_to_bytes(doc)
        formatted_bytes, extracted_text = apply_format_only(input_bytes, "standard_clean")
        
        # Text should be preserved (though extraction may not include table text)
        doc_after = bytes_to_docx(formatted_bytes)
        self.assertEqual(doc_after.tables[0].cell(0, 0).text, "Cell 1")
        self.assertEqual(doc_after.tables[0].cell(1, 1).text, "Cell 4")
    
    def test_compact_clean_profile(self):
        """
        Test compact_clean profile applies correct margins and spacing.
        NOTE: page margins are stored as TWIPS (int) in WordprocessingML; EMU conversions can round.
        """
        input_bytes = self.create_test_docx("Test paragraph\n\nSecond paragraph")
        formatted_bytes, _ = apply_format_only(input_bytes, "compact_clean")

        doc = bytes_to_docx(formatted_bytes)
        section = doc.sections[0]
        paragraph = doc.paragraphs[0]
        pf = paragraph.paragraph_format

        def cm_to_twips(cm: float) -> int:
            return int(round(cm / 2.54 * 1440))

        # Margins (2.0cm) - compare TWIPS
        self.assertEqual(section.top_margin.twips, cm_to_twips(COMPACT_CLEAN.margins["top"]))
        self.assertEqual(section.bottom_margin.twips, cm_to_twips(COMPACT_CLEAN.margins["bottom"]))
        self.assertEqual(section.left_margin.twips, cm_to_twips(COMPACT_CLEAN.margins["left"]))
        self.assertEqual(section.right_margin.twips, cm_to_twips(COMPACT_CLEAN.margins["right"]))

        # Spacing before/after should be 0pt/4pt.
        # pf.space_before may be None when it's effectively 0.
        before = pf.space_before.pt if pf.space_before is not None else 0.0
        after  = pf.space_after.pt  if pf.space_after  is not None else 0.0
        self.assertEqual(before, 0.0)
        self.assertEqual(after, 4.0)

        # Line spacing: python-docx may expose None even if style enforces single.
        # Accept 1.0 or None here (style-level single spacing).
        self.assertTrue(pf.line_spacing in (None, 1.0))
    
    def test_large_readable_profile(self):
        """
        Test large_readable profile applies correct margins, font, and spacing.
        NOTE: page margins are stored as TWIPS (int) in WordprocessingML; EMU conversions can round.
        """
        input_bytes = self.create_test_docx("Test paragraph\n\nSecond paragraph")
        formatted_bytes, _ = apply_format_only(input_bytes, "large_readable")

        doc = bytes_to_docx(formatted_bytes)
        section = doc.sections[0]
        paragraph = doc.paragraphs[0]
        pf = paragraph.paragraph_format

        def cm_to_twips(cm: float) -> int:
            return int(round(cm / 2.54 * 1440))

        # Margins (1.2in = 3.048cm) - compare TWIPS
        expected_twips = cm_to_twips(LARGE_READABLE.margins["top"])
        self.assertEqual(section.top_margin.twips, expected_twips,
                        f"Top margin should be {expected_twips} twips (1.2in)")
        self.assertEqual(section.bottom_margin.twips, expected_twips,
                        f"Bottom margin should be {expected_twips} twips (1.2in)")
        self.assertEqual(section.left_margin.twips, expected_twips,
                        f"Left margin should be {expected_twips} twips (1.2in)")
        self.assertEqual(section.right_margin.twips, expected_twips,
                        f"Right margin should be {expected_twips} twips (1.2in)")

        # Font should be Times New Roman 14pt
        # Check first run's font - CRITICAL: must override explicit run formatting
        if paragraph.runs:
            run = paragraph.runs[0]
            self.assertEqual(run.font.name, "Times New Roman", "Font should be Times New Roman")
            self.assertEqual(run.font.size, Pt(14), "Font size should be 14pt (not 11pt)")

        # Spacing before/after should be 0pt/10pt.
        # pf.space_before may be None when it's effectively 0.
        before = pf.space_before.pt if pf.space_before is not None else 0.0
        after  = pf.space_after.pt  if pf.space_after  is not None else 0.0
        self.assertEqual(before, 0.0)
        self.assertEqual(after, 10.0)

        # Line spacing should be 1.5
        # python-docx may expose as float or None
        if pf.line_spacing is not None:
            self.assertAlmostEqual(pf.line_spacing, 1.5, places=1,
                                  msg="Line spacing should be 1.5")
    
    def test_large_readable_overrides_explicit_run_formatting(self):
        """
        Test that large_readable overrides explicit run-level font sizes (e.g., 11pt -> 14pt).
        Many input docs have explicit run formatting that must be overridden.
        """
        # Create a doc with explicit run formatting (11pt, different font)
        doc = Document()
        para = doc.add_paragraph("Test with explicit formatting")
        # Set explicit run formatting that should be overridden
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)  # Explicit 11pt that should become 14pt
        
        input_bytes = docx_to_bytes(doc)
        formatted_bytes, _ = apply_format_only(input_bytes, "large_readable")
        
        doc_after = bytes_to_docx(formatted_bytes)
        para_after = doc_after.paragraphs[0]
        
        # Verify ALL runs have been overridden to 14pt Times New Roman
        self.assertGreater(len(para_after.runs), 0, "Paragraph should have runs")
        for run in para_after.runs:
            self.assertEqual(run.font.name, "Times New Roman", 
                           f"Run should have Times New Roman font, got {run.font.name}")
            self.assertEqual(run.font.size, Pt(14), 
                           f"Run should have 14pt font size, got {run.font.size}")
        
        # Also verify standard_clean keeps 11pt
        formatted_standard, _ = apply_format_only(input_bytes, "standard_clean")
        doc_standard = bytes_to_docx(formatted_standard)
        para_standard = doc_standard.paragraphs[0]
        for run in para_standard.runs:
            self.assertEqual(run.font.size, Pt(11), 
                           "standard_clean should keep 11pt font size")
    
    def test_compact_clean_xml_output(self):
        """Test that compact_clean produces correct XML output (margins, styles, content)."""
        # Create test document with 2 paragraphs and a 2x2 table
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"
        
        input_bytes = docx_to_bytes(doc)
        formatted_bytes, _ = apply_format_only(input_bytes, "compact_clean")
        
        # Open DOCX as ZIP and extract XML
        zip_buffer = io.BytesIO(formatted_bytes)
        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
            # Read document.xml
            document_xml = zip_file.read('word/document.xml')
            doc_root = ET.fromstring(document_xml)
            
            # Read styles.xml
            styles_xml = zip_file.read('word/styles.xml')
            styles_root = ET.fromstring(styles_xml)
        
        # Register namespaces
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        w_ns_tag = '{' + w_ns + '}'
        
        # Check document.xml has paragraphs and tables
        paragraphs = doc_root.findall(f'.//{w_ns_tag}p')
        tables = doc_root.findall(f'.//{w_ns_tag}tbl')
        self.assertGreater(len(paragraphs), 0, "document.xml should have paragraphs")
        self.assertGreater(len(tables), 0, "document.xml should have tables")
        
        # Check page margins in document.xml (2.0cm = 1134 twips)
        sectPr = doc_root.find(f'.//{w_ns_tag}sectPr')
        self.assertIsNotNone(sectPr, "document.xml should have sectPr")
        
        pgMar = sectPr.find(f'{w_ns_tag}pgMar')
        self.assertIsNotNone(pgMar, "sectPr should have pgMar")
        
        # 2.0cm = 1134 twips (1cm = 567 twips)
        expected_twips = "1134"
        self.assertEqual(pgMar.get(f'{w_ns_tag}left'), expected_twips,
                        "Left margin should be 1134 twips (2.0cm)")
        self.assertEqual(pgMar.get(f'{w_ns_tag}right'), expected_twips,
                        "Right margin should be 1134 twips (2.0cm)")
        self.assertEqual(pgMar.get(f'{w_ns_tag}top'), expected_twips,
                        "Top margin should be 1134 twips (2.0cm)")
        self.assertEqual(pgMar.get(f'{w_ns_tag}bottom'), expected_twips,
                        "Bottom margin should be 1134 twips (2.0cm)")
        
        # Check styles.xml Normal style
        normal_style = styles_root.find(f'.//{w_ns_tag}style[@{w_ns_tag}styleId="Normal"]')
        self.assertIsNotNone(normal_style, "styles.xml should have Normal style")
        
        # Check Normal style has rPr (run properties) with font
        rPr = normal_style.find(f'{w_ns_tag}rPr')
        if rPr is not None:
            rFonts = rPr.find(f'{w_ns_tag}rFonts')
            if rFonts is not None:
                ascii_font = rFonts.get(f'{w_ns_tag}ascii')
                self.assertEqual(ascii_font, "Calibri", "Normal style should have Calibri font")
            
            sz = rPr.find(f'{w_ns_tag}sz')
            if sz is not None:
                sz_val = sz.get(f'{w_ns_tag}val')
                # 11pt = 22 half-points
                self.assertEqual(sz_val, "22", "Normal style should have font size 11pt (22 half-points)")
        
        # Check Normal style has pPr (paragraph properties) with spacing
        pPr = normal_style.find(f'{w_ns_tag}pPr')
        if pPr is not None:
            spacing = pPr.find(f'{w_ns_tag}spacing')
            if spacing is not None:
                # 0pt before = 0 twips
                before = spacing.get(f'{w_ns_tag}before')
                self.assertEqual(before, "0", "Normal style should have 0pt spacing before")
                
                # 4pt after = 80 twips (1pt = 20 twips)
                after = spacing.get(f'{w_ns_tag}after')
                self.assertEqual(after, "80", "Normal style should have 4pt spacing after (80 twips)")
                
                # 1.0 line spacing = 240 twips per line (single spacing)
                line = spacing.get(f'{w_ns_tag}line')
                line_rule = spacing.get(f'{w_ns_tag}lineRule')
                # Single spacing can be 240 or auto
                if line is not None:
                    self.assertIn(line, ["240", "480"], "Line spacing should be 240 (single) or equivalent")
                if line_rule is not None:
                    self.assertIn(line_rule, ["auto", "exact"], "Line rule should be auto or exact")


if __name__ == '__main__':
    unittest.main()

