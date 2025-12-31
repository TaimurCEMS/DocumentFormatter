import functions_framework
from google.cloud import firestore
from google.cloud import storage
import requests
from docx import Document
from docx.shared import Pt
import io
import os
import json
import base64
import tempfile
import uuid
import logging

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Lazy initialization to avoid import-time credential errors
_db_client = None
_storage_client = None

def get_db():
    """Lazy initialization of Firestore client."""
    global _db_client
    if _db_client is None:
        _db_client = firestore.Client()
    return _db_client

def get_storage():
    """Lazy initialization of Storage client."""
    global _storage_client
    if _storage_client is None:
        _storage_client = storage.Client()
    return _storage_client

OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
BUCKET_NAME = "documentformatterapp.firebasestorage.app"


@functions_framework.cloud_event
def process_document_worker(cloud_event):
    doc_id = None
    doc_ref = None
    
    try:
        print(f"Received cloud event type: {type(cloud_event)}")
        print(f"Cloud event data type: {type(cloud_event.data)}")
        print(f"Cloud event data: {cloud_event.data}")
        
        # Handle different Pub/Sub event formats
        doc_id = None
        
        # Format 1: Standard Pub/Sub CloudEvent format
        if isinstance(cloud_event.data, dict):
            if "message" in cloud_event.data:
                message = cloud_event.data["message"]
                if "data" in message:
                    try:
                        # Decode base64 message data
                        if isinstance(message["data"], str):
                            pubsub_message = base64.b64decode(message["data"]).decode('utf-8')
                        else:
                            pubsub_message = base64.b64decode(message["data"]).decode('utf-8')
                        print(f"Decoded Pub/Sub message: {pubsub_message}")
                        message_data = json.loads(pubsub_message)
                        doc_id = message_data.get('doc_id')
                        if doc_id:
                            logger.info(f"WORKER_REV=FORMATTER_V1 doc_id={doc_id}")
                    except Exception as decode_error:
                        print(f"Error decoding message: {str(decode_error)}")
                        raise
        
        # Format 2: Direct data access (alternative format)
        if not doc_id and isinstance(cloud_event.data, dict):
            if 'doc_id' in cloud_event.data:
                doc_id = cloud_event.data.get('doc_id')
                if doc_id:
                    logger.info(f"WORKER_REV=FORMATTER_V1 doc_id={doc_id}")
            elif 'data' in cloud_event.data:
                # Try to parse data field directly
                try:
                    data_str = cloud_event.data.get('data')
                    if isinstance(data_str, str):
                        message_data = json.loads(data_str)
                        doc_id = message_data.get('doc_id')
                        if doc_id:
                            logger.info(f"WORKER_REV=FORMATTER_V1 doc_id={doc_id}")
                except:
                    pass
        
        # Format 3: Check if data is a string that needs parsing
        if not doc_id and isinstance(cloud_event.data, str):
            try:
                message_data = json.loads(cloud_event.data)
                doc_id = message_data.get('doc_id')
                if doc_id:
                    logger.info(f"WORKER_REV=FORMATTER_V1 doc_id={doc_id}")
            except:
                pass
        
        if not doc_id:
            error_msg = f"Could not extract doc_id from cloud event. Event data: {cloud_event.data}"
            print(f"Error: {error_msg}")
            raise ValueError(error_msg)
        
        print(f"Processing job with doc_id: {doc_id}")
        
        # Get clients
        db = get_db()
        storage_client = get_storage()
        
        # Use new jobs collection
        doc_ref = db.collection('jobs').document(doc_id)
        doc = doc_ref.get()
        
        if not doc.exists:
            print(f"Job {doc_id} not found")
            return
        
        data = doc.to_dict()
        
        # Check current state - ensure idempotency (don't go backwards)
        current_state = data.get('state', 'QUEUED')
        if current_state not in ('QUEUED', 'PROCESSING'):
            print(f"Job {doc_id} is already in state {current_state}, skipping")
            return
        
        storage_path = data.get('storage_path')
        style_prompt = data.get('style_prompt', 'Formal Academic Style')
        
        # Read mode and style/profile from job (defaults for backward compatibility)
        mode = data.get('mode', 'format_only')
        style = data.get('style') or data.get('profile', 'standard_clean')
        
        # Log job fields
        logger.info(f"MODE={mode} STYLE={style} STORAGE_PATH={storage_path}")
        
        if not storage_path:
            raise ValueError("Missing storage_path in job data")
        
        # Update to PROCESSING state with progress
        doc_ref.update({
            'state': 'PROCESSING',
            'status': 'PROCESSING',  # Alias (keep identical to state)
            'progress': 5,
            'display_message': 'Processing',
            'updated_at': firestore.SERVER_TIMESTAMP
        })
        
        # Step 1: Download document (progress: 20%)
        doc_ref.update({
            'progress': 20,
            'display_message': 'Downloading document',
            'updated_at': firestore.SERVER_TIMESTAMP
        })
        
        # Download the DOCX file
        # Note: Download exceptions are preserved as-is (not converted to "empty document" errors)
        obj_path = normalize_storage_path(storage_path)
        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(obj_path)
        
        if not blob.exists():
            raise ValueError(f"File not found at storage path: {storage_path}")
        
        # Download as bytes
        input_docx_bytes = blob.download_as_bytes()
        
        # Step 2: Process based on mode (progress: 50%)
        doc_ref.update({
            'progress': 50,
            'display_message': 'Formatting your document',
            'updated_at': firestore.SERVER_TIMESTAMP
        })
        
        if mode == 'format_only':
            # Format-only mode: preserve text, apply formatting
            # Note: No empty text validation for format_only - proceed even if document is empty
            from formatting.formatter_engine import apply_format_only
            logger.info("FORMAT_ONLY: apply_format_only invoked")
            try:
                output_bytes, formatted_text = apply_format_only(input_docx_bytes, style)
            except Exception as format_error:
                print(f"Error in format_only processing: {str(format_error)}")
                raise
        else:
            # Legacy mode: extract text, format text, create new DOCX
            extracted_text = download_and_extract_text(storage_path)
            
            # Only validate text content for non-format_only flows
            if not extracted_text or not extracted_text.strip():
                raise ValueError("Document is empty or contains no text")
            
            formatted_text = format_text_basic(extracted_text, style_prompt)
            
            # Generate formatted DOCX
            formatted_doc = Document()
            for line in formatted_text.split('\n'):
                if line.strip():
                    p = formatted_doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
            
            output_buffer = io.BytesIO()
            formatted_doc.save(output_buffer)
            output_buffer.seek(0)
            output_bytes = output_buffer.getvalue()
        
        # Step 3: Upload to Cloud Storage (progress: 90%)
        doc_ref.update({
            'progress': 90,
            'display_message': 'Uploading formatted document',
            'updated_at': firestore.SERVER_TIMESTAMP
        })
        
        # Use output path format: outputs/{docId}_formatted.docx
        object_path = f'outputs/{doc_id}_formatted.docx'
        token = str(uuid.uuid4())
        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(object_path)
        
        # Upload using upload_from_string (NOT upload_from_file)
        blob.upload_from_string(
            output_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Set metadata with download token after upload
        blob.metadata = {"firebaseStorageDownloadTokens": token}
        
        # Ensure metadata persisted
        blob.patch()
        
        # Verify metadata was set (reload blob)
        blob.reload()
        
        # Construct Firebase download URL (NO signed URLs - Firebase token URL only)
        from urllib.parse import quote
        encoded = quote(object_path, safe="")
        download_url = f"https://firebasestorage.googleapis.com/v0/b/{BUCKET_NAME}/o/{encoded}?alt=media&token={token}"
        
        # RUNTIME TRIPWIRE: Detect any signed URLs (forbidden)
        if "generate_signed" in str(download_url) or "X-Goog-Signature" in str(download_url):
            raise RuntimeError("SIGNED URL DETECTED - forbidden")
        
        # Log download URL for debugging
        print(f"DOWNLOAD_URL= {download_url}")
        
        # Update to COMPLETED state with all required fields including download_url
        doc_ref.update({
            'state': 'COMPLETED',
            'status': 'COMPLETED',  # Alias (keep identical to state)
            'progress': 100,
            'display_message': 'Completed',
            'formatted_text': formatted_text,
            'download_url': download_url,  # FlutterFlow uses this directly
            'error': None,
            'updated_at': firestore.SERVER_TIMESTAMP
        })
        
        print(f"Job {doc_id} completed successfully")
        
    except Exception as e:
        error_msg = str(e)
        import traceback
        full_traceback = traceback.format_exc()
        
        print(f"ERROR processing job {doc_id if doc_id else 'unknown'}: {error_msg}")
        print(f"Full traceback:\n{full_traceback}")
        
        # Try to update Firestore with error, even if doc_id/doc_ref aren't set
        if doc_id:
            try:
                if doc_ref:
                    # Check current state to avoid going backwards
                    try:
                        current_doc = doc_ref.get()
                        if current_doc.exists:
                            current_data = current_doc.to_dict()
                            current_state = current_data.get('state', 'QUEUED')
                            # Only update if not already in a terminal state
                            if current_state not in ('COMPLETED', 'FAILED'):
                                doc_ref.update({
                                    'state': 'FAILED',
                                    'status': 'FAILED',  # Alias (keep identical to state)
                                    'display_message': 'Failed',
                                    'error': error_msg,
                                    'updated_at': firestore.SERVER_TIMESTAMP
                                })
                    except:
                        # Fallback: try to update anyway
                        doc_ref.update({
                            'state': 'FAILED',
                            'status': 'FAILED',  # Alias (keep identical to state)
                            'display_message': 'Failed',
                            'error': error_msg,
                            'updated_at': firestore.SERVER_TIMESTAMP
                        })
                else:
                    # Create doc_ref if we have doc_id but doc_ref wasn't created
                    doc_ref = db.collection('jobs').document(doc_id)
                    doc_ref.update({
                        'state': 'FAILED',
                        'status': 'FAILED',  # Alias (keep identical to state)
                        'display_message': 'Failed',
                        'error': error_msg,
                        'updated_at': firestore.SERVER_TIMESTAMP
                    })
                print(f"Updated Firestore with error for doc_id: {doc_id}")
            except Exception as update_error:
                print(f"CRITICAL: Failed to update job status in Firestore: {str(update_error)}")
                print(f"Update error traceback:\n{traceback.format_exc()}")
        else:
            print(f"CRITICAL: Cannot update Firestore - doc_id is None. Error: {error_msg}")
            print(f"This means the error occurred before doc_id could be extracted from the Pub/Sub message.")


def normalize_storage_path(value):
    """
    Normalize storage path to object path for blob operations.
    - If starts with "http": extract after "/o/" before "?", URL-decode -> object path
    - If starts with "gs://": strip "gs://<bucket>/" -> object path
    - Else: return as-is
    """
    if not value:
        return value
    
    # Handle HTTP/HTTPS URLs (Firebase Storage URLs)
    if value.startswith('http://') or value.startswith('https://'):
        try:
            from urllib.parse import unquote, urlparse, parse_qs
            # Extract path after "/o/"
            if '/o/' in value:
                # Get the part after "/o/"
                parts = value.split('/o/', 1)
                if len(parts) == 2:
                    # Get path before query string
                    path_part = parts[1].split('?')[0]
                    # URL decode
                    obj_path = unquote(path_part)
                    return obj_path
            # Fallback: try to parse as URL
            parsed = urlparse(value)
            if '/o/' in parsed.path:
                path_part = parsed.path.split('/o/', 1)[1]
                obj_path = unquote(path_part)
                return obj_path
            # If no /o/ found, return as-is (might be direct download URL)
            return value
        except Exception as e:
            print(f"Warning: Could not normalize HTTP URL {value}: {e}")
            return value
    
    # Handle gs:// URLs
    elif value.startswith('gs://'):
        # Strip "gs://<bucket>/" -> object path
        parts = value.replace('gs://', '').split('/', 1)
        if len(parts) == 2:
            # Return just the object path (skip bucket name)
            return parts[1]
        else:
            # Invalid format, return as-is
            return value
    
    # Else: return as-is
    return value


def download_and_extract_text(storage_path):
    """Download .docx file and extract text. Supports gs:// and Firebase HTTPS URLs."""
    try:
        # Handle gs:// URLs - use storage client
        if storage_path.startswith('gs://'):
            storage_client = get_storage()
            # Parse gs:// URL
            parts = storage_path.replace('gs://', '').split('/', 1)
            if len(parts) != 2:
                raise ValueError(f"Invalid gs:// URL format: {storage_path}")
            bucket_name = parts[0]
            blob_name = parts[1]
            
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(blob_name)
            
            if not blob.exists():
                raise ValueError(f"File not found at path: {blob_name}")
            
            # Download as bytes
            file_bytes = blob.download_as_bytes()
            
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
                tmp.write(file_bytes)
                tmp.flush()
            
            # Extract text using python-docx
            doc = Document(tmp_path)
            text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Clean up temp file
            os.unlink(tmp_path)
            return text
        
        # Handle Firebase HTTPS URLs - use requests
        elif storage_path.startswith('http://') or storage_path.startswith('https://'):
            # Download file as binary using requests
            response = requests.get(storage_path, stream=True)
            response.raise_for_status()
            
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
                for chunk in response.iter_content(chunk_size=8192):
                    tmp.write(chunk)
                tmp.flush()
            
            # Extract text using python-docx
            doc = Document(tmp_path)
            text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Clean up temp file
            os.unlink(tmp_path)
            return text
        
        else:
            # Assume it's a relative path in default bucket
            storage_client = get_storage()
            bucket = storage_client.bucket(BUCKET_NAME)
            blob = bucket.blob(storage_path)
            
            if not blob.exists():
                raise ValueError(f"File not found at path: {storage_path}")
            
            # Download as bytes
            file_bytes = blob.download_as_bytes()
            
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
                tmp.write(file_bytes)
                tmp.flush()
            
            # Extract text using python-docx
            doc = Document(tmp_path)
            text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Clean up temp file
            os.unlink(tmp_path)
            return text
                
    except Exception as e:
        print(f"Error downloading/extracting text from {storage_path}: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return None


def format_text_basic(text, style_prompt):
    """
    Basic V1 formatting rules (no OpenAI dependency).
    Applies simple formatting: proper paragraph spacing, sentence capitalization.
    """
    if not text or not text.strip():
        return text
    
    # Split into paragraphs
    paragraphs = text.split('\n\n')
    formatted_paragraphs = []
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # Ensure sentences start with capital letter
        sentences = para.split('. ')
        formatted_sentences = []
        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Capitalize first letter of sentence
            if sentence and not sentence[0].isupper():
                sentence = sentence[0].upper() + sentence[1:] if len(sentence) > 1 else sentence.upper()
            
            # Add period if missing (except last sentence if it already has one)
            if i < len(sentences) - 1 and not sentence.endswith('.'):
                sentence += '.'
            
            formatted_sentences.append(sentence)
        
        formatted_para = '. '.join(formatted_sentences)
        formatted_paragraphs.append(formatted_para)
    
    # Join paragraphs with double newline
    result = '\n\n'.join(formatted_paragraphs)
    
    # Ensure text ends with period if it doesn't
    if result and not result.rstrip().endswith(('.', '!', '?')):
        result = result.rstrip() + '.'
    
    return result


def format_with_openai(text, style_prompt):
    """OpenAI formatting (kept for future use, not used in V1)."""
    headers = {
        'Authorization': f'Bearer {OPENAI_API_KEY}',
        'Content-Type': 'application/json'
    }
    
    payload = {
        'model': 'gpt-4o-mini',
        'messages': [
            {'role': 'system', 'content': f'You are a document formatter. Apply {style_prompt} formatting to the text. Return only the formatted text, no explanations.'},
            {'role': 'user', 'content': text}
        ],
        'temperature': 0.3
    }
    
    response = requests.post('https://api.openai.com/v1/chat/completions', headers=headers, json=payload )
    response.raise_for_status()
    
    return response.json()['choices'][0]['message']['content']
